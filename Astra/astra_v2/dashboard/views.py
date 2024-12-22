from django.shortcuts import render, get_object_or_404, redirect, get_list_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.db.models import Q
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.contrib import messages

from .models import Question, Document as DocModel, Tag
from .forms import QuestionForm, AnswerFormSet, DocumentForm

from sandbox.sandbox import Sandbox
from ids_core import IDS
from alert_bot.telegram_bot import send_telegram_alert, send_daily_report

import json
import zipfile
import os
from docx import Document as DocxDocument

# Инициализация IDS
ids = IDS()

def check_access_or_redirect(request):
    """
    Проверяет доступ пользователя через context_processors.py.
    Если доступа нет, перенаправляет на главную страницу.
    """
    if not request.context.get('has_service_access', False):
        messages.error(request, "У вас нет доступа к этому разделу.")
        return redirect('home')

@login_required
def dashboard(request):
    check_access_or_redirect(request)
    return render(request, 'dashboard/dashboard.html', {
        'current_page': 'dashboard'
    })

@login_required
def questions(request):
    check_access_or_redirect(request)
    query = request.GET.get('q')
    if query:
        questions = Question.objects.filter(
            Q(text__icontains=query) | Q(tags__name__icontains=query)
        ).distinct()
    else:
        questions = Question.objects.all()
    return render(request, 'dashboard/questions.html', {
        'questions': questions,
        'current_page': 'questions'
    })

@login_required
def documents(request):
    check_access_or_redirect(request)
    query = request.GET.get('q')
    selected_ids = request.GET.get('selected_ids', '').split(',')
    if query:
        documents = Document.objects.filter(
            Q(title__icontains=query) | Q(tags__name__icontains=query)
        ).distinct()
    else:
        documents = Document.objects.all()
    return render(request, 'dashboard/documents.html', {
        'documents': documents,
        'selected_ids': selected_ids,
        'current_page': 'documents'
    })

@login_required
def create_question(request):
    check_access_or_redirect(request)
    if request.method == 'POST':
        tags = request.POST.getlist('tags')
        new_tags = []
        for tag in tags:
            if not tag.isdigit():
                tag_obj, created = Tag.objects.get_or_create(name=tag)
                new_tags.append(str(tag_obj.id))
            else:
                new_tags.append(tag)

        request.POST = request.POST.copy()
        request.POST.setlist('tags', new_tags)

        form = QuestionForm(request.POST)
        formset = AnswerFormSet(request.POST)

        if form.is_valid() and formset.is_valid():
            question = form.save()
            formset.instance = question
            formset.save()
            ids.log_event(request.user.username, "Создание нового вопроса")
            return redirect('dashboard')
    else:
        form = QuestionForm()
        formset = AnswerFormSet()

    return render(request, 'dashboard/create_question.html', {
        'form': form,
        'formset': formset,
        'current_page': 'create_question'
    })

@login_required
def upload_document(request):
    check_access_or_redirect(request)
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            ids.log_event(request.user.username, "Загрузка документа")
            return redirect('documents')
    else:
        form = DocumentForm()
    return render(request, 'dashboard/upload_document.html', {
        'form': form,
        'current_page': 'upload_document'
    })


# # Инициализация IDS с отправкой уведомлений
# ids = IDS(alert_func=send_telegram_alert, report_func=send_daily_report)

# def check_user_group(user):
#     return user.groups.filter(name__in=['Администратор', 'Методист']).exists()

# @login_required
# @user_passes_test(check_user_group)
# def dashboard(request):
#     query = request.GET.get('q')
#     if query:
#         questions = Question.objects.filter(
#             Q(text__icontains=query) | Q(tags__name__icontains=query)
#         ).distinct()
#     else:
#         questions = Question.objects.all()

#     documents = DocModel.objects.all()
#     return render(request, 'dashboard/dashboard.html', {'questions': questions, 'documents': documents, 'current_page': 'dashboard'})

# @login_required
# @user_passes_test(check_user_group)
# def questions(request):
#     query = request.GET.get('q')
#     if query:
#         questions = Question.objects.filter(
#             Q(text__icontains=query) | Q(tags__name__icontains=query)
#         ).distinct()
#     else:
#         questions = Question.objects.all()
#     return render(request, 'dashboard/questions.html', {'questions': questions, 'current_page': 'questions'})

# @login_required
# @user_passes_test(check_user_group)
# def documents(request):
#     query = request.GET.get('q')
#     selected_ids = request.GET.get('selected_ids', '').split(',')
#     if query:
#         documents = DocModel.objects.filter(
#             Q(title__icontains=query) | Q(tags__name__icontains=query)
#         ).distinct()
#     else:
#         documents = DocModel.objects.all()
#     return render(request, 'dashboard/documents.html', {
#         'documents': documents,
#         'selected_ids': selected_ids,
#         'current_page': 'documents'
#     })

# @login_required
# @user_passes_test(check_user_group)
# def edit_question(request, pk):
#     question = get_object_or_404(Question, pk=pk)
#     if request.method == 'POST':
#         tags = request.POST.getlist('tags')
#         new_tags = []
#         for tag in tags:
#             if not tag.isdigit():
#                 tag_obj, created = Tag.objects.get_or_create(name=tag)
#                 new_tags.append(str(tag_obj.id))
#             else:
#                 new_tags.append(tag)

#         request.POST = request.POST.copy()
#         request.POST.setlist('tags', new_tags)

#         form = QuestionForm(request.POST, instance=question)
#         formset = AnswerFormSet(request.POST, instance=question)

#         if (form.is_valid() and formset.is_valid()):
#             question = form.save()
#             formset.save()
#             ids.log_event(request.user.username, "Редактирование вопроса")  # Логирование действия редактирования вопроса
#             return redirect('dashboard')
#     else:
#         form = QuestionForm(instance=question)
#         formset = AnswerFormSet(instance=question)

#     tags = list(question.tags.values_list('id', flat=True))
#     return render(request, 'dashboard/edit_question.html', {'form': form, 'formset': formset, 'question': question, 'tags': tags, 'current_page': 'edit_question'})

# @login_required
# @user_passes_test(check_user_group)
# def create_question(request):
#     if request.method == 'POST':
#         tags = request.POST.getlist('tags')
#         new_tags = []
#         for tag in tags:
#             if not tag.isdigit():
#                 tag_obj, created = Tag.objects.get_or_create(name=tag)
#                 new_tags.append(str(tag_obj.id))
#             else:
#                 new_tags.append(tag)

#         request.POST = request.POST.copy()
#         request.POST.setlist('tags', new_tags)

#         form = QuestionForm(request.POST)
#         formset = AnswerFormSet(request.POST)

#         if form.is_valid() and formset.is_valid():
#             question = form.save()
#             formset.instance = question
#             formset.save()
#             ids.log_event(request.user.username, "Создание нового вопроса")  # Логирование создания нового вопроса
#             return redirect('dashboard')
#     else:
#         form = QuestionForm()
#         formset = AnswerFormSet()

#     return render(request, 'dashboard/create_question.html', {'form': form, 'formset': form, 'current_page': 'create_question'})

# @login_required
# @user_passes_test(check_user_group)
# def upload_document(request):
#     # Инициализируем IDS и передаем его в Sandbox
#     sandbox = Sandbox(ids, user=request.user.username)
#     print("[DEBUG] Sandbox инициализирован с IDS")
#     print("[DEBUG] Начало выполнения upload_document")

#     if request.method == 'POST':
#         print(f"[DEBUG] request.FILES содержимое: {request.FILES}")

#         form = DocumentForm(request.POST, request.FILES)
#         if form.is_valid() and 'file' in request.FILES:
#             uploaded_file = request.FILES['file']
#             print(f"[DEBUG] Файл {uploaded_file.name} получен для загрузки")

#             # Сохраняем файл во временной папке для проверки
#             temp_path = sandbox.save_temp_file(uploaded_file)

#             if temp_path is None:
#                 print("[ERROR] Ошибка сохранения временного файла")
#                 messages.error(request, "Ошибка при сохранении временного файла для проверки.")
#                 return render(request, 'dashboard/upload_document.html', {'form': form, 'current_page': 'upload_document'})

#             # Запускаем процесс проверки файла, передавая путь к файлу
#             result = sandbox.process_file(temp_path)
#             print(f"[DEBUG] Результат проверки файла: {result}")

#             # Проверка на безопасность файла
#             if result != "Файл успешно прошел все проверки.":
#                 ids.log_event(request.user.username, f"Файл {uploaded_file.name} не прошел проверку безопасности", critical=True)
#                 if os.path.exists(temp_path):
#                     os.remove(temp_path)
#                 print("[ERROR] Файл не прошел проверку, возврат на страницу загрузки")
#                 messages.error(request, result)
#                 return render(request, 'dashboard/upload_document.html', {'form': form, 'current_page': 'upload_document'})

#             # Если проверка пройдена, сохраняем файл в media/documents
#             final_path = f"documents/{uploaded_file.name}"
#             saved_path = default_storage.save(final_path, ContentFile(uploaded_file.read()))

#             # Ручное создание записи документа с указанием пути к файлу
#             document = form.save(commit=False)
#             document.file.name = saved_path
#             document.save()

#             # Удаление временного файла
#             if os.path.exists(temp_path):
#                 os.remove(temp_path)

#             print("[DEBUG] Документ успешно загружен в media")
#             ids.log_event(request.user.username, f"Файл {uploaded_file.name} успешно загружен", critical=False)
#             messages.success(request, "Документ успешно загружен.")
#             return redirect('dashboard')

#     else:
#         form = DocumentForm()

#     print("[DEBUG] Открытие страницы загрузки документа")
#     return render(request, 'dashboard/upload_document.html', {'form': form, 'current_page': 'upload_document'})






# @login_required
# @user_passes_test(check_user_group)
# def edit_document(request, pk):
#     document = get_object_or_404(DocModel, pk=pk)
#     if request.method == 'POST':
#         form = DocumentForm(request.POST, request.FILES, instance=document)
#         if form.is_valid():
#             form.save()
#             ids.log_event(request.user.username, "Редактирование документа")  # Логирование редактирования документа
#             return redirect('dashboard')
#     else:
#         form = DocumentForm(instance=document)
    
#     return render(request, 'dashboard/edit_document.html', {'form': form, 'document': document, 'current_page': 'edit_document'})

# @login_required
# @user_passes_test(check_user_group)
# def view_document(request, pk):
#     document = get_object_or_404(DocModel, pk=pk)
#     ids.log_event(request.user.username, "Просмотр документа")  # Логирование просмотра документа
#     return render(request, 'dashboard/view_document.html', {'document': document, 'current_page': 'view_document'})

# @login_required
# @user_passes_test(check_user_group)
# @require_POST
# def export_questions(request):
#     data = json.loads(request.body)
#     question_ids = data.get('ids', [])
#     questions = get_list_or_404(Question, id__in=question_ids)
    
#     doc = DocxDocument()
#     for question in questions:
#         doc.add_heading(f'Вопрос {question.id}', level=1)
#         doc.add_paragraph(question.text)
#         doc.add_heading('Ответы:', level=2)
#         for answer in question.answers.all():
#             answer_paragraph = doc.add_paragraph()
#             run = answer_paragraph.add_run(answer.text)
#             if answer.correct:
#                 run.bold = True

#     response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     response['Content-Disposition'] = 'attachment; filename=selected_questions.docx'
#     doc.save(response)
#     ids.log_event(request.user.username, "Экспорт вопросов")  # Логирование экспорта вопросов
#     return response

# @login_required
# @user_passes_test(check_user_group)
# @require_POST
# def export_documents(request):
#     data = json.loads(request.body)
#     document_ids = data.get('ids', [])
#     documents = get_list_or_404(DocModel, id__in=document_ids)
    
#     zip_filename = "selected_documents.zip"
#     response = HttpResponse(content_type='application/zip')
#     response['Content-Disposition'] = f'attachment; filename={zip_filename}'

#     with zipfile.ZipFile(response, 'w') as zip_file:
#         for document in documents:
#             file_path = document.file.path
#             zip_file.write(file_path, os.path.basename(file_path))

#     ids.log_event(request.user.username, "Экспорт документов")  # Логирование экспорта документов
#     return response

# @login_required
# @user_passes_test(check_user_group)
# def tag_list(request):
#     tags = Tag.objects.all()
#     return render(request, 'dashboard/tags.html', {'tags': tags, 'current_page': 'tags'})

# @login_required
# @user_passes_test(check_user_group)
# def tag_autocomplete(request):
#     if 'q' in request.GET:
#         query = request.GET.get('q')
#         tags = Tag.objects.filter(name__icontains=query).values('id', 'name')
#         results = [{'id': tag['id'], 'text': tag['name']} for tag in tags]
#         return JsonResponse(results, safe=False)
#     return JsonResponse([], safe=False)

# @login_required
# @user_passes_test(check_user_group)
# def preview_document(request, pk):
#     document = get_object_or_404(DocModel, pk=pk)
#     ids.log_event(request.user.username, "Предпросмотр документа")  # Логирование предпросмотра документа
#     return render(request, 'dashboard/preview_document.html', {'document': document})

# def generate_report(request):
#     # Запуск генерации отчета
#     ids.generate_daily_report()
#     return HttpResponse("Отчет успешно создан")